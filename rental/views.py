from django.http import HttpResponse, JsonResponse
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.forms import AuthenticationForm
from django.contrib.auth import login,logout
from django.urls import reverse_lazy
from django.views.generic import View,UpdateView,DeleteView
from django.contrib.auth.mixins import LoginRequiredMixin
from django.contrib.auth.hashers import make_password
from docx import Document
from django.views import View
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from .models import Prokat
from django.contrib.auth.views import LoginView, LogoutView
from .forms import IshchiForm, QurilmaForm, TamirlashForm, UserLoginForm, MijozForm, ProkotForm, QaytarishForm, \
    CustomUserChangeForm
from django.db.models import Count
import openpyxl
from openpyxl.styles import Alignment, Font, Border, Side
from .models import CustomUser, Mijozlar, Qaytarish,Qurilmalar,Prokat, Tamirlash


class ProfileEditView(UpdateView):
    model = CustomUser
    form_class = CustomUserChangeForm
    template_name = 'edit_profile.html'
    success_url = reverse_lazy('home')

class Loginview(LoginRequiredMixin,View):
    def get(self,request):
        userlogin=UserLoginForm()
        return render(request,'registration/login.html',{'form':userlogin})
    
    def post(self,request):
        check=UserLoginForm(data=request.POST)
        if check.is_valid():
            user=check.get_user()
            login(request,user)
          
            return render(request,'home.html')
        else:
            return redirect('login')
class LogoutView(LoginRequiredMixin,View):
         def get(self,request):
             logout(request)
             return redirect('login')
class HomeView(LoginRequiredMixin,View):
    def get(self,request):
        mijozlar_soni=Mijozlar.objects.count()
        jami_qurilmalar = Qurilmalar.objects.count()
        faol_shartnomalar = Prokat.objects.filter(qaytarish__isnull=True).count()
        qaytarilgan_soni = Qaytarish.objects.count()
        tamirlashdagi_qurilmalar = Tamirlash.objects.count()

        # context orqali ma'lumotlarni templatega uzatamiz
        context = {
            'mijozlar_soni':mijozlar_soni,
            'jami_qurilmalar': jami_qurilmalar,
            'faol_shartnomalar': faol_shartnomalar,
            'qaytarilgan_soni': qaytarilgan_soni,
            'tamirlashdagi_qurilmalar': tamirlashdagi_qurilmalar,
        }
        return render(request,'home.html',context)


from django.contrib import messages

class EditProfileView(LoginRequiredMixin,UpdateView):
    model = CustomUser
    form_class = CustomUserChangeForm
    template_name = 'edit_profile.html'
    success_url = reverse_lazy('home')

class MijozView(LoginRequiredMixin,View):
    def get(self,request):
        mijozlar=Mijozlar.objects.all()
        return render(request,'mijoz_qushish.html',{'mijoz':mijozlar})
    

class MijozAddView(LoginRequiredMixin,View):
    def get(self,request):
        mijozlar=Mijozlar.objects.all()
        mijoz=MijozForm()
        return render(request,'mijoz_qushish.html',{'form':mijoz,'mijoz':mijozlar})
    def post(self,request):
        mijozlar=Mijozlar.objects.all()
        mijoz=MijozForm(data=request.POST)
        if mijoz.is_valid():
            mijoz.save()
            return redirect('mijoz')
        else:
            return redirect('mijoz',{'message':'Malumot saqlanmadi','mijoz':mijozlar})

class MijozDeleteView(LoginRequiredMixin,View):
    def get(self,request,pk):
        mijoz=Mijozlar.objects.get(id=pk)
        mijoz.delete()
        return redirect('mijoz')
class MijozUpdateView(LoginRequiredMixin,UpdateView):
    model = Mijozlar
    form_class = MijozForm
    template_name = 'edit_mijoz.html'
    success_url = reverse_lazy('mijoz')

class QurilmaView(LoginRequiredMixin,View):
    def get(self,request):
        qurilma=Qurilmalar.objects.all()
        return render(request,'qurilmalar.html',{'qurilmalar':qurilma})
    
class QurilmaQushView(LoginRequiredMixin,View):
    def get(self, request):
        qurilma_form = QurilmaForm()
        qurilmalar = Qurilmalar.objects.all()
        return render(request, 'qurilmalar.html', {'form': qurilma_form, 'qurilmalar': qurilmalar})

    def post(self, request):
        qurilma_form = QurilmaForm(data=request.POST,files=request.FILES)
        qurilmalar = Qurilmalar.objects.all()
        if qurilma_form.is_valid():
            qurilma_form.save()
            return redirect('qurilmalar')
        else:
            return render(request, 'qurilmalar.html', {'form': qurilma_form, 'qurilmalar': qurilmalar, 'message': 'Qurilma saqlanmadi'})

class QurilmaUpdateView(LoginRequiredMixin,UpdateView):
    model = Qurilmalar
    form_class=QurilmaForm
    template_name = 'qurilma_edit.html'
    success_url = reverse_lazy('qurilmalar')

class QurilmaDeleteView(LoginRequiredMixin,View):
    def get(self, request, pk):
        qurilma = Qurilmalar.objects.get(id=pk)
        qurilma.delete()
        return redirect('qurilmalar')

class ShartnomaAllView(LoginRequiredMixin,View):
    def get(self,request):
        shartnomalar = Prokat.objects.all()
        return render(request, 'shartnoma.html', {'shartnomalar': shartnomalar})
    
class ShartnomaView(LoginRequiredMixin,View):
    def get(self, request):
        shartnoma = ProkotForm()
        shartnomalar = Prokat.objects.all()
        return render(request, 'shartnoma.html', {'form': shartnoma, 'shartnomalar': shartnomalar})

    def post(self, request):
        shartnoma = ProkotForm(data=request.POST)
        shartnomalar = Prokat.objects.all()
        if shartnoma.is_valid():
            shartnoma.save()
            message = 'Shartnoma muaffaqiyatli tuzildi'
            return redirect("prokatall")
        else:
            message = "Shartnoma saqlanmadi"
        
        return render(request, 'shartnoma.html', {'form': shartnoma, 'shartnomalar': shartnomalar, 'message': message})
    
class ShartnomaEditView(LoginRequiredMixin,UpdateView):
    model=Prokat
    form_class=ProkotForm
    template_name='edit_prokot.html'
    success_url=reverse_lazy('prokot')
    
class ShartnomaDeleteView(LoginRequiredMixin,View):
     def get(self, request, pk):
        shartnoma = Prokat.objects.get(id=pk)
        shartnoma.delete()
        return redirect('prokatall')
    
    


class ShartnomaDownloadView(LoginRequiredMixin,View):
    def get(self, request, pk):
        shartnoma = Prokat.objects.get(id=pk)
        document = Document()

        # Heading
        heading = document.add_heading(level=0)
        run = heading.add_run('Хорезмская область прокатная компания ИНН 665899871416, 620102, г.\n'
                              'Аль Хорезми, ул. Телефон: +998(99)-999-99-99')
        run.font.size = Pt(13)
        run.bold = True
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        title = document.add_paragraph('Квитанция об оплате')
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Order details
        table = document.add_table(rows=3, cols=2)
        table.cell(0, 0).text = 'Номер заказа:'
        table.cell(0, 1).text = f'{shartnoma.id:08d}'
        table.cell(1, 0).text = 'Дата начала:'
        table.cell(1, 1).text = shartnoma.berilgan_sanasi.strftime('%d.%m.%Y %H:%M')
        table.cell(2, 0).text = 'Дата окончания:'
        table.cell(2, 1).text = shartnoma.qaytish_sanasi.strftime('%d.%m.%Y %H:%M')

        document.add_paragraph(f'ФИО заказчика: {shartnoma.mijoz.ism}')

        
        # Item details
        document.add_paragraph()
        table = document.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Название'
        hdr_cells[1].text = 'Залог (сум)'
        hdr_cells[2].text = 'Цена проката за сутки (сум)'
        hdr_cells[3].text = 'Количество'

        row_cells = table.add_row().cells
        row_cells[0].text = shartnoma.qurilma.nomi
        row_cells[1].text = str(shartnoma.avans)
        row_cells[2].text = str(shartnoma.kunlik_narxi)
        row_cells[3].text = str(shartnoma.miqdori)

        # Adding borders to the table
        for row in table.rows:
            for cell in row.cells:
                tc_pr = cell._element.get_or_add_tcPr()
                tc_borders = OxmlElement('w:tcBorders')
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'single')
                    tc_borders.append(border)
                tc_pr.append(tc_borders)

        # Summary details
        document.add_paragraph()
        document.add_paragraph(f'Сумма залога (сум): {shartnoma.avans}')
        document.add_paragraph(f'Стоимость проката с учётом количества дней (сум): {shartnoma.kunlik_narxi}')
        document.add_paragraph(f'Количество дней: {shartnoma.kunlar_soni}')
        document.add_paragraph()
        document.add_paragraph(f'Сумма к оплате: {shartnoma.umumiy_narx} (сум)')
        document.add_paragraph()

        # Signatures
        document.add_paragraph('Подпись заказчика: _____________________________________')
        document.add_paragraph(f'{shartnoma.mijoz.ism}')
        document.add_paragraph('Подпись сотрудника: ____________________________________')
        document.add_paragraph(f'{shartnoma.ishchi.ism}')

        # Response
        file_name = f'shartnoma_{shartnoma.id}.docx'
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename={file_name}'
        document.save(response)

        return response




class ExportAllContractsExcelView(LoginRequiredMixin,View):
    def get(self, request, *args, **kwargs):
        # Prokat modelidan barcha yozuvlarni olish
        contracts = Prokat.objects.all()

        # Excel fayl yaratish
        workbook = openpyxl.Workbook()
        sheet = workbook.active
        sheet.title = "Barcha Shartnomalar"

        # Chegaralar, shriftlar va joylashuv
        bold_font = Font(bold=True, size=12)
        regular_font = Font(size=12)
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        alignment = Alignment(horizontal="center", vertical="center")

        # Sarlavhalar
        headers = [
            "ID", "MIJOZ ISMI","MIJOZ FAMILIYASI", "USKUNA", "Miqdori", "MASTER",
            "BERILGAN SANA", "QAYTARISH SANASI",
            "Kunlar soni", "Montaj",
            "Xizmat narxi", "Avans",
            "Kunlik narxi", "Umumiy narx"
        ]

        # Sarlavhalarni yozish
        for col_num, header in enumerate(headers, 1):
            cell = sheet.cell(row=1, column=col_num)
            cell.value = header
            cell.font = bold_font
            cell.border = thin_border
            cell.alignment = alignment

        # Ma'lumotlarni yozish
        for row_num, contract in enumerate(contracts, start=2):
            row_data = [
                contract.id,
                contract.mijoz.ism,
                contract.mijoz.familiya,
                contract.qurilma.nomi,
                contract.miqdori,
                contract.ishchi.ism,
                contract.berilgan_sanasi.strftime("%Y-%m-%d") if contract.berilgan_sanasi else "—",
                contract.qaytish_sanasi.strftime("%Y-%m-%d") if contract.qaytish_sanasi else "—",
                contract.kunlar_soni,
                "Ha" if contract.montaj else "Yo'q",
                contract.xizmat_narxi,
                contract.avans,
                contract.kunlik_narxi,
                contract.umumiy_narx
            ]
            for col_num, value in enumerate(row_data, 1):
                cell = sheet.cell(row=row_num, column=col_num)
                cell.value = value
                cell.font = regular_font
                cell.border = thin_border
                cell.alignment = alignment

        # Ustun kengliklarini sozlash
        column_widths = [5, 20, 25, 10, 20, 15, 15, 12, 10, 15, 15, 15, 15]
        for i, width in enumerate(column_widths, 1):
            sheet.column_dimensions[openpyxl.utils.get_column_letter(i)].width = width

        # Excel faylni javobga biriktirish
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        response['Content-Disposition'] = 'attachment; filename="barcha_shartnomalar.xlsx"'
        workbook.save(response)
        return response


    


class QaytarishView(LoginRequiredMixin, View):
    def get(self, request):
        form = QaytarishForm()
        malumot = Qaytarish.objects.all()
        malumot_id = Qaytarish.objects.values_list('prokat', flat=True)
        form.fields['prokat'].queryset = Prokat.objects.exclude(id__in=malumot_id)
        return render(request, 'qaytarish.html', {'form': form, 'malumotlar': malumot})

    def post(self, request):
        malumot = Qaytarish.objects.all()
        form = QaytarishForm(data=request.POST)
        malumot_id = Qaytarish.objects.values_list('prokat', flat=True)
        form.fields['prokat'].queryset = Prokat.objects.exclude(id__in=malumot_id) 
        if form.is_valid():
            form.save()
            return redirect('qaytarish')
        else:
            message = 'Saqlanmadi: ' + str(form.errors) 
            return render(request, 'qaytarish.html', {'form': form, 'malumotlar': malumot, 'message': message})


class QaytarishAllView(LoginRequiredMixin,View):
    def get(self,request):
            malumot=Qaytarish.objects.all()
            return render(request,'qaytarish.html',{'malumotlar':malumot})
        
class QaytarishDeleteView(LoginRequiredMixin,View):
    def get(self,request,pk):
        malumot=Qaytarish.objects.get(id=pk)
        malumot.delete()
        return redirect ('qaytarish')
class QaytarishEditView(LoginRequiredMixin,UpdateView):
    model = Qaytarish
    form_class=QaytarishForm
    template_name = 'qaytarish_edit.html'
    success_url = reverse_lazy('qaytarish')
    
    
    


class QaytarishDocxDownloadView(LoginRequiredMixin,View):
    def get(self, request, pk):
        qaytarish = Qaytarish.objects.get(id=pk)
        document = Document()

        # Heading
        heading = document.add_heading(level=0)
        run = heading.add_run('Хорезмская область прокатная компания ИНН 665899871416, 620102, г.\n'
                              'Аль Хорезми, ул. Телефон: +998(99)-999-99-99')
        run.font.size = Pt(13)  # Setting smaller font size
        run.bold = True
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        title = document.add_paragraph('Квитанция о возврате')
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Order details
        table = document.add_table(rows=3, cols=2)
        table.cell(0, 0).text = 'Номер заказа:'
        table.cell(0, 1).text = f'{qaytarish.prokat.id:08d}'
        table.cell(1, 0).text = 'Дата начала:'
        table.cell(1, 1).text = qaytarish.prokat.berilgan_sanasi.strftime('%d.%m.%Y %H:%M')
        table.cell(2, 0).text = 'Дата возврата:'
        table.cell(2, 1).text = qaytarish.qaytarilgan_sana.strftime('%d.%m.%Y %H:%M')

        document.add_paragraph(f'ФИО заказчика: {qaytarish.prokat.mijoz.ism}')

        # Adding borders to the table
        

        # Item details
        document.add_paragraph()
        table = document.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Название'
        hdr_cells[1].text = 'Залог (сум)'
        hdr_cells[2].text = 'Цена проката за сутки (сум)'
        hdr_cells[3].text = 'Количество'

        row_cells = table.add_row().cells
        row_cells[0].text = qaytarish.prokat.qurilma.nomi
        row_cells[1].text = str(qaytarish.prokat.avans)
        row_cells[2].text = str(qaytarish.prokat.kunlik_narxi)
        row_cells[3].text = str(qaytarish.prokat.miqdori)

        # Adding borders to the table
        for row in table.rows:
            for cell in row.cells:
                tc_pr = cell._element.get_or_add_tcPr()
                tc_borders = OxmlElement('w:tcBorders')
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'single')
                    tc_borders.append(border)
                tc_pr.append(tc_borders)

        # Summary details
        document.add_paragraph()
        document.add_paragraph(f'Сумма залога (сум): {qaytarish.prokat.avans}')
        document.add_paragraph(f'Стоимость проката с учётом количества дней (сум): {qaytarish.prokat.kunlik_narxi}')
        document.add_paragraph(f'Количество дней: {qaytarish.prokat.kunlar_soni}')
        document.add_paragraph()
        document.add_paragraph(f'Сумма к оплате: {qaytarish.prokat.umumiy_narx} (сум)')
        document.add_paragraph()

        # Signatures
        document.add_paragraph('Подпись заказчика: _____________________________________')
        document.add_paragraph(f'{qaytarish.prokat.mijoz.ism}')
        document.add_paragraph('Подпись сотрудника: ____________________________________')
        document.add_paragraph(f'{qaytarish.prokat.ishchi.ism}')

        # Response
        file_name = f'qaytarish_{qaytarish.id}.docx'
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename={file_name}'
        document.save(response)

        return response

    
class TamirlashView(LoginRequiredMixin, View):
    def get(self, request):
        tamirlash = Tamirlash.objects.all()
        return render(request, 'tamirlash.html', {'tamirlash': tamirlash})

class TamirlashQoshView(LoginRequiredMixin,View):
    def get(self, request):
        tamirlash_form = TamirlashForm()
        tamirlash = Tamirlash.objects.all()
        tamirlash_qurilmalari = Tamirlash.objects.values_list('qurilma_id', flat=True)
        tamirlash_form.fields['qurilma'].queryset = Qurilmalar.objects.exclude(id__in=tamirlash_qurilmalari)
        return render(request, 'tamirlash.html', {'form': tamirlash_form,'tamirlash': tamirlash})
        
    def post(self, request):
        tamirlash_form = TamirlashForm(request.POST)
        tamirlash = Tamirlash.objects.all()
        if tamirlash_form.is_valid():
            qurilma = tamirlash_form.save()
            return redirect('tamirlash')  
        return render(request, 'tamirlash.html', {'form': tamirlash_form,'tamirlash': tamirlash})
class TamirlashEditView(LoginRequiredMixin, UpdateView):
    model = Tamirlash
    form_class = TamirlashForm
    template_name = 'tamirlash_edit.html'
    success_url = reverse_lazy('tamirlash')


class TamirlashDeleteView(LoginRequiredMixin, View):
    def get(self, request, pk):
        tamirlash = Tamirlash.objects.get(id=pk)
        tamirlash.delete()
        return redirect('tamirlash')

class IshchiView(LoginRequiredMixin, View):
    def get(self, request):
        ishchi = CustomUser.objects.all()
        return render(request, 'ishchi.html', {'ishchi': ishchi})

from django.shortcuts import render, redirect
from django.contrib.auth.hashers import make_password
from django.views import View
from django.contrib.auth.mixins import LoginRequiredMixin
from .forms import IshchiForm
from .models import CustomUser

class IshchiQoshishView(LoginRequiredMixin, View):
    def get(self, request):
        ishchi_form = IshchiForm()
        ishchilar = CustomUser.objects.all()
        return render(request, 'ishchi.html', {'form': ishchi_form, 'ishchi': ishchilar})

    def post(self, request):
        ishchi_form = IshchiForm(data=request.POST)
        ishchilar = CustomUser.objects.all()
        if ishchi_form.is_valid():
            new_ishchi = ishchi_form.save(commit=False)
            new_ishchi.password = make_password(ishchi_form.cleaned_data['password'])
            new_ishchi.save()
            return redirect('ishchilar')  
        else:
            return render(request, 'ishchi.html', {'form': ishchi_form, 'ishchi': ishchilar, 'message': 'Ishchi saqlanmadi'})

class IshchiEditView(LoginRequiredMixin, UpdateView):
    model = CustomUser
    form_class = IshchiForm
    template_name = 'ishchi_edit.html'
    success_url = reverse_lazy('ishchilar')

    def get_object(self, queryset=None):
        return get_object_or_404(CustomUser, id=self.kwargs['id'])


class IshchiDeleteView(LoginRequiredMixin, View):
    def get(self, request, id):
        ischi = CustomUser.objects.get(id=id)
        ischi.delete()
        return redirect('ishchilar')